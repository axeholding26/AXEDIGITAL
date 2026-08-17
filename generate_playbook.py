"""
Génère le playbook de projet Axe Digital au format .docx
Auteur : Axe Digital
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Couleurs de la charte Axe Digital
BRAND_BLUE = RGBColor(0x02, 0x38, 0xD6)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Styles de base ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def set_cell_background(cell, hex_color):
    """Applique une couleur de fond à une cellule."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading_custom(text, level=1, color=BRAND_BLUE, size=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return h

def add_paragraph(text, bold_prefix=None, color=None, size=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        if color:
            r.font.color.rgb = color
    r = p.add_run(text)
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def add_table(headers, rows, header_bg='0238D6'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        set_cell_background(hdr[i], header_bg)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AXE DIGITAL')
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = BRAND_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PLAYBOOK DE PROJET')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('& État des lieux complet')
run.font.size = Pt(18)
run.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cabinet de Consulting IA & Solutions Numériques pour l\'Afrique')
run.font.size = Pt(13)
run.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0  •  Douala, Cameroun')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES MANUELLE
# ════════════════════════════════════════════════════════════
add_heading_custom('Sommaire', level=1, color=DARK, size=22)
toc = [
    ('1.', 'Présentation du projet'),
    ('2.', 'État des lieux technique'),
    ('3.', 'Architecture & Structure du code'),
    ('4.', 'Contenu & Sections du site'),
    ('5.', 'Playbook de gestion de projet'),
    ('6.', 'Conventions de code & Bonnes pratiques'),
    ('7.', 'Workflow Git & Déploiement'),
    ('8.', 'Feuille de route & Prochaines étapes'),
    ('9.', 'Informations pratiques & Contact'),
]
for num, title in toc:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(13)
    run.font.color.rgb = BRAND_BLUE
    run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. PRÉSENTATION DU PROJET
# ════════════════════════════════════════════════════════════
add_heading_custom('1. Présentation du projet', level=1, size=24)

add_heading_custom('1.1 Vue d\'ensemble', level=2, size=16)
add_paragraph(
    'Axe Digital est un cabinet de consulting en Intelligence Artificielle et développement de '
    'solutions numériques, spécialisé dans l\'accompagnement des entreprises africaines vers leur '
    'transformation digitale. Le projet consiste en une landing page haute performance (Single Page '
    'Application) conçue pour présenter l\'offre, générer des leads et convertir les visiteurs via '
    'WhatsApp et les formules tarifaires.'
)

add_heading_custom('1.2 Vision', level=2, size=16)
add_paragraph(
    'Propulser les organisations africaines dans l\'ère du digital grâce à des agents IA, des '
    'applications métiers et des solutions numériques sur-mesure, pensées et conçues pour le '
    'contexte local (connexions instables, usage mobile, réalités du marché africain).'
)

add_heading_custom('1.3 Objectifs du site', level=2, size=16)
add_bullet('Présenter l\'offre de services (agents IA, applications métiers, applications mobiles).')
add_bullet('Générer des leads qualifiés via des CTA WhatsApp et des appels découverte.')
add_bullet('Vendre les formules tarifaires (Découverte, Croissance, Performance).')
add_bullet('Instaurer la confiance via les témoignages, réalisations et statistiques.')
add_bullet('Optimiser le SEO local (Cameroun / Afrique francophone).')

add_heading_custom('1.4 Positionnement', level=2, size=16)
add_paragraph(
    'Différenciation forte : solutions « conçues pour l\'Afrique », jamais importées. Expertise '
    'centrale en agents IA, applications offline-first, maintenance locale incluse et accompagnement '
    'de A à Z, face à une approche traditionnelle ou à des freelances/agences classiques.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. ÉTAT DES LIEUX TECHNIQUE
# ════════════════════════════════════════════════════════════
add_heading_custom('2. État des lieux technique', level=1, size=24)

add_heading_custom('2.1 Stack technologique', level=2, size=16)
add_table(
    ['Catégorie', 'Technologie', 'Rôle'],
    [
        ['Framework', 'React 19.2', 'Interface utilisateur (SPA)'],
        ['Langage', 'TypeScript 5.8', 'Typage statique'],
        ['Build', 'Vite 6.2', 'Bundler & dev server (port 3000)'],
        ['Styling', 'Tailwind CSS (CDN)', 'Design system & utilitaires'],
        ['Animations', 'Framer Motion 12', 'Animations UI'],
        ['Animations avancées', 'GSAP 3.15', 'Animations agents robotiques'],
        ['Icônes', 'lucide-react 0.554', 'Icônes vectorielles'],
        ['Compteurs', 'react-countup 6.5', 'Animations de statistiques'],
        ['Observateur', 'react-intersection-observer 10', 'Déclenchement au scroll'],
        ['3D', 'model-viewer', 'Visualisation agents 3D (GLB)'],
    ]
)

add_heading_custom('2.2 Dépendances installées', level=2, size=16)
add_paragraph('Dépendances (production) :')
add_bullet('framer-motion ^12.23.24')
add_bullet('gsap ^3.15.0')
add_bullet('lucide-react ^0.554.0')
add_bullet('react ^19.2.0')
add_bullet('react-countup ^6.5.3')
add_bullet('react-dom ^19.2.0')
add_bullet('react-intersection-observer ^10.0.0')
add_paragraph('Dépendances (développement) :')
add_bullet('@types/node ^22.14.0')
add_bullet('@vitejs/plugin-react ^5.0.0')
add_bullet('typescript ~5.8.2')
add_bullet('vite ^6.2.0')

add_heading_custom('2.3 Fonctionnalités du site', level=2, size=16)
add_bullet('Preloader animé avec logo de chargement (preload.png).')
add_bullet('Halo suiveur du curseur (CursorGlow), désactivé sur tactile.')
add_bullet('Navbar fixe avec menu mobile et navigation fluide (scroll).')
add_bullet('Hero avec globe animé de bulles (canvas) et bandeau de logos partenaires défilant.')
add_bullet('Compteurs statistiques animés (react-countup).')
add_bullet('Portfolio vidéo (autoplay, loop, muted) avec overlay au survol.')
add_bullet('Section « Pourquoi ça stagne » (3 problèmes).')
add_bullet('Présentation des 3 solutions sur-mesure.')
add_bullet('Formules tarifaires avec lien de paiement externe (mychariow.shop).')
add_bullet('Comparatif Axe Digital vs alternatives.')
add_bullet('Témoignages clients (6 avis) et statistiques de satisfaction.')
add_bullet('CTA final avec appel découverte et rendez-vous.')
add_bullet('Footer avec coordonnées, réseaux sociaux et liens WhatsApp.')
add_bullet('Intégration Meta Pixel pour le suivi des conversions.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. ARCHITECTURE & STRUCTURE DU CODE
# ════════════════════════════════════════════════════════════
add_heading_custom('3. Architecture & Structure du code', level=1, size=24)

add_heading_custom('3.1 Arborescence', level=2, size=16)
structure = (
    'axe-digital/\n'
    '├─ App.tsx                    # Composant racine qui orchestre les sections\n'
    '├─ index.html                 # HTML de base (SEO, Meta Pixel, Tailwind CDN)\n'
    '├─ index.tsx                  # Point d\'entrée React\n'
    '├─ package.json               # Dépendances & scripts\n'
    '├─ vite.config.ts             # Configuration Vite (port 3000, alias @)\n'
    '├─ tsconfig.json              # Configuration TypeScript\n'
    '├─ metadata.json              # Métadonnées du projet\n'
    '├─ README.md                  # Documentation d\'installation\n'
    '├─ TODO.md                    # Suivi des tâches\n'
    '├─ public/\n'
    '│  ├─ assets/                 # Images, logos, vidéos, modèles 3D\n'
    '│  ├─ robots.txt              # SEO robots\n'
    '│  └─ sitemap.xml             # Sitemap SEO\n'
    '└─ src/\n'
    '   ├─ global.d.ts             # Types globaux (model-viewer)\n'
    '   ├─ images.d.ts             # Déclarations d\'imports d\'images\n'
    '   ├─ image/AppImage.tsx      # Registre central des assets\n'
    '   ├─ utils/whatsapp.ts       # Utilitaires WhatsApp\n'
    '   └─ components/             # Composants de sections'
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = DARK

add_heading_custom('3.2 Composants (src/components)', level=2, size=16)
components = [
    'Preloader.tsx', 'CursorGlow.tsx', 'Navbar.tsx', 'Hero.tsx', 'HeroGlobe.tsx',
    'Stats.tsx', 'Portfolio.tsx', 'ProblemSection.tsx', 'SolutionsSection.tsx',
    'PricingSection.tsx', 'WhySection.tsx', 'SolutionSection.tsx',
    'TestimonialSection.tsx', 'CTASection.tsx', 'Footer.tsx',
    'AIAgent.tsx', 'AIAgentsContainer.tsx', 'AIAgent.css',
]
for c in components:
    add_bullet(c)

add_heading_custom('3.3 Ordre des sections dans App.tsx', level=2, size=16)
flow = (
    'Preloader → CursorGlow → Navbar → Hero → Stats → Portfolio → ProblemSection → '
    'SolutionsSection → PricingSection → WhySection → SolutionSection → '
    'TestimonialSection → CTASection → Footer'
)
add_paragraph(flow)

add_heading_custom('3.4 Registre des assets (AppImage)', level=2, size=16)
add_paragraph(
    'Toutes les images du site sont référencées de manière centralisée dans '
    'src/image/AppImage.tsx (classe statique). Cela permet une maintenance simple '
    'des chemins d\'assets. Les chemins sont relatifs au dossier public/ (ex: /assets/Logo/logo.png).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. CONTENU & SECTIONS DU SITE
# ════════════════════════════════════════════════════════════
add_heading_custom('4. Contenu & Sections du site', level=1, size=24)

add_heading_custom('4.1 Navigation (Navbar)', level=2, size=16)
nav_links = [
    'Nos réalisations → #nos-realisations',
    'Nos solutions → #solutions-ia',
    'Nos formules → #nos-formules',
    'Ils nous font confiance → #ils-nous-font-confiance',
    'Nous contacter → #nous-contacter',
]
for l in nav_links:
    add_bullet(l)
add_paragraph('CTA principal : « Demander un audit » (ouvre WhatsApp).')

add_heading_custom('4.2 Les 3 solutions', level=2, size=16)
add_table(
    ['Solution', 'Description'],
    [
        ['Agents IA intelligents', 'Assistant virtuels 24/7 : service client, qualification de leads, traitement de documents.'],
        ['Applications métiers sur-mesure', 'CRM, ERP, RH, stock, comptabilité adaptés au secteur et au marché africain.'],
        ['Applications mobiles pour l\'Afrique', 'Android/iOS offline-first, légères, adaptées aux connexions instables.'],
    ]
)

add_heading_custom('4.3 Formules tarifaires', level=2, size=16)
add_table(
    ['Pack', 'Prix (FCFA)', 'Public / Contenu'],
    [
        ['Découverte', '350 000', 'Audit SI, identification des processus automatisables, recommandations IA, rapport + 1 séance.'],
        ['Croissance', '750 000', 'Tout Découverte + agent IA pilote, app métier simple, suivi & optimisation, support prioritaire.'],
        ['Performance', '1 500 000', 'Tout Croissance + agents IA à grande échelle, apps complexes, app mobile, équipe dédiée, maintenance.'],
    ]
)
add_paragraph('Lien de paiement externe commun : https://rnlwqugt.mychariow.shop/starter85')

add_heading_custom('4.4 Réalisations (Portfolio)', level=2, size=16)
add_table(
    ['Projet', 'Catégorie'],
    [
        ['FINITRANS', 'Application de gestion de Transit et Dédouanement'],
        ['Stella AI', 'Agent intelligent pour PC'],
        ['SULTAN', 'Application de gestion d\'une boucherie moderne'],
    ]
)

add_heading_custom('4.5 Témoignages clients', level=2, size=16)
add_table(
    ['Client', 'Entreprise', 'Résultat clé'],
    [
        ['Richkard Kamena', 'Lyxera', '80% du support automatisé, revenus doublés en 3 mois.'],
        ['Junior Tenevock', 'Zaninii', 'Application métier adaptée, gestion opérationnelle simplifiée.'],
        ['Eleanor Pena', 'Mapa Sarl', 'Taux de conversion augmenté de 40%.'],
        ['Jalil Ketou', 'Skynesis', 'Meilleure positionnement et acquisition constante de clients.'],
        ['Marcel Kop', 'Symphoni Social', 'Disponibilité, réactivité, véritable partenariat.'],
        ['Merveille Noumo', 'La grande Royale', 'Optimisation des coûts via digitalisation pas à pas.'],
    ]
)

add_heading_custom('4.6 Statistiques affichées', level=2, size=16)
add_bullet('+50 projets IA et solutions numériques réalisés depuis 2024.')
add_bullet('+250% de gain de productivité moyen pour les clients.')
add_bullet('+97% de clients satisfaits.')
add_bullet('100% clients satisfaits, note moyenne 5.0, +50 projets, 2x croissance moyenne.')
add_bullet('ROI +350% en 6 mois (section SolutionSection).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. PLAYBOOK DE GESTION DE PROJET
# ════════════════════════════════════════════════════════════
add_heading_custom('5. Playbook de gestion de projet', level=1, size=24)

add_heading_custom('5.1 Cycle de vie d\'un projet', level=2, size=16)
phases = [
    ('1. Audit & Découverte', 'Analyse du système d\'information, identification des processus automatisables, recommandations stratégiques IA.'),
    ('2. Cadrage & Proposition', 'Définition du périmètre, choix de la formule (Découverte/Croissance/Performance), proposition commerciale.'),
    ('3. Conception (Design)', 'Maquettage UI/UX, architecture des solutions, validation du client.'),
    ('4. Développement', 'Développement des agents IA, applications métiers et/ou mobiles, itérations.'),
    ('5. Test & Validation', 'Tests fonctionnels, optimisation, recette client.'),
    ('6. Déploiement & Formation', 'Mise en production, accompagnement pas à pas, formation des équipes.'),
    ('7. Suivi & Optimisation', 'Maintenance, support prioritaire, évolution continue.'),
]
for t, d in phases:
    add_bullet(d, bold_prefix=f'{t} : ')

add_heading_custom('5.2 Rôles & Responsabilités', level=2, size=16)
add_table(
    ['Rôle', 'Responsabilités'],
    [
        ['Chef de projet', 'Coordination, planning, interface client, suivi des livrables.'],
        ['Développeur', 'Développement front-end/back-end, intégration des solutions.'],
        ['Data Scientist', 'Conception et entraînement des agents IA.'],
        ['Designer UI/UX', 'Maquettes, design system, expérience utilisateur.'],
        ['Consultant IA', 'Stratégie, audit, recommandations, accompagnement.'],
    ]
)

add_heading_custom('5.3 Processus de livraison', level=2, size=16)
add_bullet('Chaque fonctionnalité est développée en itération courte et validée par le client.')
add_bullet('Les livrables sont présentés avec une démonstration et une documentation.')
add_bullet('Un rapport de suivi est fourni à chaque phase.')
add_bullet('La communication se fait principalement via WhatsApp (réactivité 7j/7).')

add_heading_custom('5.4 Gestion des priorités (TODO.md)', level=2, size=16)
add_paragraph(
    'Le projet utilise un fichier TODO.md à la racine pour suivre les tâches en français. '
    'Toute tâche est cochée ([x]) une fois réalisée et vérifiée. Il est recommandé de regrouper '
    'les tâches par domaine (design, fonctionnalité, SEO, build).'
)

add_heading_custom('5.5 Indicateurs de succès (KPI)', level=2, size=16)
add_bullet('Nombre de leads générés (clics WhatsApp, appels découverte).')
add_bullet('Taux de conversion des formules.')
add_bullet('Performance de chargement (Lighthouse, Core Web Vitals).')
add_bullet('Positionnement SEO local (Google Cameroun).')
add_bullet('Satisfaction client (note moyenne cible : 5.0).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. CONVENTIONS DE CODE
# ════════════════════════════════════════════════════════════
add_heading_custom('6. Conventions de code & Bonnes pratiques', level=1, size=24)

add_heading_custom('6.1 Syntaxe & Style', level=2, size=16)
add_bullet('Composants React écrits en TypeScript avec typed props (React.FC).')
add_bullet('Nommage en PascalCase pour les composants, camelCase pour les fonctions/variables.')
add_bullet('Utilisation des déclarations de modules d\'images (src/images.d.ts).')
add_bullet('Classes Tailwind utilitaires + custom CSS localisé quand nécessaire (style JSX).')

add_heading_custom('6.2 Charte graphique', level=2, size=16)
add_table(
    ['Élément', 'Valeur'],
    [
        ['Couleur principale (brand-blue)', '#0238D6'],
        ['Couleur noire (textes/CTA)', '#1A1A1A / #000000'],
        ['Fond de section clair', '#f8fafc (slate-50)'],
        ['Police', 'Poppins (300 à 800)'],
        ['Arrière-plan général', '#ffffff'],
    ]
)

add_heading_custom('6.3 Bonnes pratiques', level=2, size=16)
add_bullet('Centraliser les images dans AppImage.tsx pour éviter les chemins dupliqués.')
add_bullet('Utiliser les messages WhatsApp pré-définis dans src/utils/whatsapp.ts.')
add_bullet('Ajouter les identifiants de section (id) pour la navigation fluide.')
add_bullet('Vérifier le build (npm run build) avant chaque livraison.')
add_bullet('Respecter les stratégies d\'animation (framer-motion) et les états au scroll.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. WORKFLOW GIT & DÉPLOIEMENT
# ════════════════════════════════════════════════════════════
add_heading_custom('7. Workflow Git & Déploiement', level=1, size=24)

add_heading_custom('7.1 Commandes principales', level=2, size=16)
cmds = [
    ('npm install', 'Installer les dépendances'),
    ('npm run dev', 'Lancer le serveur de développement (port 3000)'),
    ('npm run build', 'Compiler le projet pour la production (dist/)'),
    ('npm run preview', 'Prévisualiser le build de production'),
]
add_table(['Commande', 'Description'], cmds)

add_heading_custom('7.2 Workflow Git recommandé', level=2, size=16)
add_bullet('Utiliser des branches par fonctionnalité (ex: feature/nouvelle-section).')
add_bullet('Commits clairs et descriptifs en français ou en anglais.')
add_bullet('Fusionner dans la branche principale (main) après validation du build.')
add_bullet('Documenter les changements dans TODO.md et README.md.')

add_heading_custom('7.3 Déploiement', level=2, size=16)
add_paragraph(
    'Le site est un build Vite statique. Après npm run build, le contenu du dossier dist/ '
    'est prêt à être hébergé (Netlify, Vercel, Serveur, etc.). Le domaine prévu est '
    'https://axedigital.agency. Le fichier .gitignore exclut node_modules, dist et les fichiers .local.'
)

add_heading_custom('7.4 SEO & Analytics', level=2, size=16)
add_bullet('Meta tags SEO dans index.html (description, keywords, Open Graph, Twitter).')
add_bullet('Données géographiques (Douala, Cameroun, coordonnées GPS).')
add_bullet('robots.txt et sitemap.xml présents dans public/.')
add_bullet('Meta Pixel (Facebook) intégré pour le suivi des conversions.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. FEUILLE DE ROUTE
# ════════════════════════════════════════════════════════════
add_heading_custom('8. Feuille de route & Prochaines étapes', level=1, size=24)

add_heading_custom('8.1 État actuel (fait)', level=2, size=16)
completed = [
    'Exploration des fichiers (Footer.tsx, Hero.tsx) et définition du plan.',
    'Footer : ajout de la localisation « Ange Raphael, fin barrière ESSEC ».',
    'Hero : réduction de la hauteur du logo Kreads (alignement avec les autres logos).',
    'Vérification du build/lint.',
    'Implémentation de l\'effet cercle suiveur sur le curseur (CursorGlow).',
    'Création du preloader avec preload.png comme logo (Preloader.tsx).',
    'Intégration du Preloader dans App.tsx.',
]
for t in completed:
    add_bullet(t)

add_heading_custom('8.2 Prochaines étapes recommandées', level=2, size=16)
recommended = [
    'Mettre en place un routage de page (si multi-pages) ou étendre les sections.',
    'Optimiser les performances (chargement des vidéos, lazy-loading, images).',
    'Continuer l\'optimisation SEO et le raffinement des méta-données.',
    'Ajouter une section FAQ / blog pour améliorer le référencement.',
    'Centraliser les configurations (couleurs, polices) dans un fichier de design system.',
    'Automatiser le déploiement (CI/CD) sur l\'hébergeur cible.',
    'Configurer les variables d\'environnement (.env.local) pour les clés API.',
    'Réaliser des tests de performance et d\'accessibilité (Lighthouse, axe).',
]
for t in recommended:
    add_bullet(t)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9. INFORMATIONS PRATIQUES
# ════════════════════════════════════════════════════════════
add_heading_custom('9. Informations pratiques & Contact', level=1, size=24)

add_heading_custom('9.1 Coordonnées', level=2, size=16)
add_table(
    ['Élément', 'Valeur'],
    [
        ['Entreprise', 'Axe Digital'],
        ['Secteur', 'Consulting IA & Solutions numériques'],
        ['Localisation', 'Ange Raphael, fin barrière ESSEC — Douala, Cameroun'],
        ['Email', 'contact@axedigital.agency'],
        ['Téléphone', '(+237) 656523837'],
        ['Téléphone WhatsApp', '(+237) 671715511'],
        ['Site web', 'https://axedigital.agency'],
        ['LinkedIn', 'linkedin.com/company/axe-digital-community'],
        ['Facebook', 'facebook.com (Axe Digital)'],
    ]
)

add_heading_custom('9.2 Messages WhatsApp pré-définis', level=2, size=16)
add_bullet('Commander : « Bonjour, je souhaite commander vos services... »')
add_bullet('Appel : « Bonjour, je souhaite passer un appel pour discuter de vos services. »')
add_bullet('Contact : « Bonjour, je souhaite obtenir plus d\'informations sur vos services. »')
add_bullet('Pack : « Bonjour, je suis intéressé(e) par le pack [nom]. Pouvez-vous me donner plus de détails ? »')

add_heading_custom('9.3 Réseaux sociaux (Footer)', level=2, size=16)
add_bullet('LinkedIn & Facebook (actifs).')
add_bullet('WhatsApp (bouton direct).')
add_bullet('TikTok (lien présent, à compléter).')

add_heading_custom('9.4 Conclusion', level=2, size=16)
add_paragraph(
    'Ce playbook centralise la connaissance du projet Axe Digital pour faciliter la montée en '
    'compétence des équipes, standardiser les processus et servir de référence de gestion. '
    'Il doit être mis à jour à chaque étape clé du projet.'
)

# ════════════════════════════════════════════════════════════
# Signature
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Axe Digital • Playbook de projet v1.0 —')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

# ── En-tête et pied de page ──
section = doc.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'AXE DIGITAL — Playbook de projet'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

footer = section.footer
fp = footer.paragraphs[0]
fp.text = 'Cabinet de Consulting IA & Solutions Numériques adaptées à l\'Afrique'
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in fp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

# ── Sauvegarde ──
output_path = '/home/nathan/Documents/TOUT/AXE_DIGITAL/axe-digital/playbookaxedigital.docx'
doc.save(output_path)
print(f'✅ Playbook généré avec succès : {output_path}')
